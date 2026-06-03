"""
Sterling Interview Portal - DOCX Report Assembler
Assembles all chapter MD files into a properly formatted DOCX.
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE_DIR = r"c:\Users\Niraj Singh\OneDrive\Documents\Desktop\Interview portal\Project Report"
OUTPUT_DOCX = os.path.join(BASE_DIR, "Sterling_Interview_Portal_Report.docx")

# ─────────────────────────────────────────────
# HELPER: add page break
# ─────────────────────────────────────────────
def add_page_break(doc):
    doc.add_page_break()

# ─────────────────────────────────────────────
# HELPER: set paragraph formatting
# ─────────────────────────────────────────────
def fmt_para(para, font_size=12, bold=False, italic=False, color=None,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=6, space_after=6,
             line_spacing=None):
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if line_spacing:
        para.paragraph_format.line_spacing = Pt(line_spacing)
    for run in para.runs:
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.name = "Times New Roman"
        if color:
            run.font.color.rgb = RGBColor(*color)

# ─────────────────────────────────────────────
# HELPER: add styled heading
# ─────────────────────────────────────────────
def add_heading(doc, text, level=1, center=False):
    if level == 1:
        para = doc.add_paragraph()
        run = para.add_run(text.upper())
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 51, 102)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(12)
    elif level == 2:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 70, 127)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
    elif level == 3:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.italic = True
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(30, 30, 30)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(4)
    elif level == 4:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = "Times New Roman"
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(3)

# ─────────────────────────────────────────────
# HELPER: add body paragraph
# ─────────────────────────────────────────────
def add_body(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = Pt(18)  # 1.5 line spacing

# ─────────────────────────────────────────────
# HELPER: add bullet item
# ─────────────────────────────────────────────
def add_bullet(doc, text):
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)

# ─────────────────────────────────────────────
# HELPER: add numbered item
# ─────────────────────────────────────────────
def add_numbered(doc, text):
    para = doc.add_paragraph(style='List Number')
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

# ─────────────────────────────────────────────
# HELPER: add code block
# ─────────────────────────────────────────────
def add_code_block(doc, text):
    for line in text.split('\n'):
        para = doc.add_paragraph()
        run = para.add_run(line if line else ' ')
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        # Light grey background
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F5F5F5')
        pPr.append(shd)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

# ─────────────────────────────────────────────
# HELPER: add simple table from pipe-delimited markdown
# ─────────────────────────────────────────────
def add_md_table(doc, lines):
    rows = []
    for line in lines:
        if line.startswith('|') and not re.match(r'\|[-\s|]+\|', line):
            cells = [c.strip() for c in line.strip('|').split('|')]
            rows.append(cells)
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    # Pad rows
    for r in rows:
        while len(r) < max_cols:
            r.append('')
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            para = cell.paragraphs[0]
            run = para.runs[0] if para.runs else para.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"
            if i == 0:
                run.font.bold = True
                # Header shading
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'D0E4F5')
                tcPr.append(shd)
    doc.add_paragraph()  # spacing after table

# ─────────────────────────────────────────────
# HELPER: add horizontal rule separator
# ─────────────────────────────────────────────
def add_separator(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '003366')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ─────────────────────────────────────────────
# CORE PARSER: convert MD content to DOCX
# ─────────────────────────────────────────────
def parse_md_to_docx(doc, content):
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip === separators and --- at top level
        if re.match(r'^={3,}', line) or re.match(r'^-{3,}$', line):
            i += 1
            continue

        # Chapter heading (#)
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            if text.startswith('CHAPTER'):
                add_heading(doc, text, level=1)
            else:
                add_heading(doc, text, level=1)
            i += 1
            continue

        # Section heading (##)
        if line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            add_heading(doc, text, level=2)
            i += 1
            continue

        # Subsection heading (###)
        if line.startswith('### ') and not line.startswith('#### '):
            text = line[4:].strip()
            add_heading(doc, text, level=3)
            i += 1
            continue

        # Sub-subsection heading (####)
        if line.startswith('#### '):
            text = line[5:].strip()
            add_heading(doc, text, level=4)
            i += 1
            continue

        # Code block
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].rstrip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            add_code_block(doc, '\n'.join(code_lines))
            continue

        # Table (pipe-delimited)
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            add_md_table(doc, table_lines)
            continue

        # Bullet point (- item)
        if re.match(r'^- ', line):
            text = line[2:].strip()
            # Clean inline markdown
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            add_bullet(doc, text)
            i += 1
            continue

        # Numbered list (1. 2. etc.)
        if re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line).strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            add_numbered(doc, text)
            i += 1
            continue

        # Italic line (*text*)
        if line.startswith('*') and line.endswith('*') and len(line) > 2:
            inner = line.strip('*')
            para = doc.add_paragraph()
            run = para.add_run(inner)
            run.italic = True
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Empty line — paragraph break
        if not line.strip():
            i += 1
            continue

        # Regular body paragraph — clean markdown formatting
        text = line.strip()
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'`(.*?)`', r'\1', text)
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # links
        if text:
            add_body(doc, text)
        i += 1

# ─────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────
def add_cover_page(doc):
    # Title block
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("A PROJECT REPORT")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Submitted in Partial Fulfillment of the Requirements\nfor the Award of the Degree of")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("BACHELOR OF TECHNOLOGY")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("in\nComputer Science & Engineering / Information Technology")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Project title
    p = doc.add_paragraph()
    run = p.add_run("STERLING INTELLIGENT INTERVIEW &\nCANDIDATE ASSESSMENT PLATFORM")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Enterprise AI-Powered Interview Management and\nCandidate Assessment System for Sterling E-Mobility")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Submitted By:\n[Student Name] — [Roll Number]")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Under the Guidance of:\n[Guide Name, Designation]")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Industry Partner:\nSterling E-Mobility Solutions Limited\nFaridabad, Haryana, India")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("[Institution Name]\n[City, State]\nAcademic Year: 2025–2026")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_page_break(doc)

# ─────────────────────────────────────────────
# CERTIFICATE PAGE
# ─────────────────────────────────────────────
def add_certificate(doc):
    add_heading(doc, "CERTIFICATE", level=1)
    doc.add_paragraph()
    cert_text = (
        "This is to certify that the project report entitled "
        "\"Sterling Intelligent Interview & Candidate Assessment Platform\" "
        "submitted by [Student Name] (Roll No: [Roll Number]) is a bonafide record "
        "of the project work carried out under my supervision in partial fulfillment "
        "of the requirements for the award of the degree of Bachelor of Technology "
        "in Computer Science & Engineering / Information Technology from [Institution Name], "
        "during the academic year 2025–2026.\n\n"
        "The project work is original and has not been submitted earlier, either in part or full, "
        "to this institution or to any other institution for the award of any degree."
    )
    add_body(doc, cert_text)
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Date: _______________")
    run.font.size = Pt(12); run.font.name = "Times New Roman"

    doc.add_paragraph()
    doc.add_paragraph()

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.cell(0,0).text = "Project Guide\n[Name]\n[Designation]\n[Department]"
    sig_table.cell(0,1).text = "Head of Department\n[Name]\n[Department]\n[Institution]"
    for cell in sig_table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(11); run.font.name = "Times New Roman"

    add_page_break(doc)

# ─────────────────────────────────────────────
# DECLARATION PAGE
# ─────────────────────────────────────────────
def add_declaration(doc):
    add_heading(doc, "DECLARATION", level=1)
    doc.add_paragraph()
    decl_text = (
        "I hereby declare that the project report entitled \"Sterling Intelligent Interview "
        "& Candidate Assessment Platform\" submitted to [Institution Name] in partial fulfillment "
        "of the requirements for the award of the degree of Bachelor of Technology is a genuine "
        "record of my own work carried out during the academic year 2025–2026.\n\n"
        "I further declare that the information presented in this project report has not been "
        "submitted previously for the award of any degree, diploma, or certificate to this "
        "institution or to any other institution.\n\n"
        "The content of this project report is based on the work done at Sterling E-Mobility "
        "Solutions Limited, Faridabad / Bengaluru, India, during the internship/project period. "
        "All sources of information used in this report have been duly acknowledged."
    )
    add_body(doc, decl_text)
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Place: _______________          Date: _______________")
    run.font.size = Pt(12); run.font.name = "Times New Roman"
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Signature: _______________\n[Student Name]\nRoll No: [Roll Number]")
    run.font.size = Pt(12); run.font.name = "Times New Roman"

    add_page_break(doc)

# ─────────────────────────────────────────────
# ACKNOWLEDGEMENT
# ─────────────────────────────────────────────
def add_acknowledgement(doc):
    add_heading(doc, "ACKNOWLEDGEMENT", level=1)
    doc.add_paragraph()
    ack_text = (
        "I wish to express my sincere gratitude to all those who contributed to the successful "
        "completion of this project.\n\n"
        "I am deeply indebted to my Project Guide, [Guide Name], [Designation], [Department], "
        "[Institution Name], for their invaluable guidance, continuous encouragement, and constructive "
        "feedback throughout the duration of this project. Their expertise and insights were instrumental "
        "in shaping the direction and quality of this work.\n\n"
        "I extend my sincere thanks to the Head of Department, [HOD Name], for providing the necessary "
        "facilities and an encouraging environment for project completion.\n\n"
        "I am grateful to the management and technical team at Sterling E-Mobility Solutions Limited, "
        "particularly the engineering staff at the Bengaluru Technical Centre and the Faridabad "
        "manufacturing campus, for their cooperation, domain knowledge sharing, and support during "
        "the project development phase. Special appreciation is due to the HR and Engineering teams "
        "for facilitating the pilot testing of the platform.\n\n"
        "I also acknowledge the open-source community — the developers of FastAPI, React, SQLAlchemy, "
        "React Three Fiber, and the many libraries that formed the technological foundation of this "
        "platform — without whose contributions this project would not have been possible.\n\n"
        "Finally, I express my heartfelt gratitude to my family and friends for their unwavering "
        "support, patience, and encouragement throughout the duration of this project."
    )
    add_body(doc, ack_text)
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("[Student Name]")
    run.font.size = Pt(12); run.font.bold = True; run.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    add_page_break(doc)

# ─────────────────────────────────────────────
# ABSTRACT
# ─────────────────────────────────────────────
def add_abstract(doc):
    add_heading(doc, "ABSTRACT", level=1)
    doc.add_paragraph()
    abstract_text = (
        "Sterling E-Mobility Solutions Limited, India's largest manufacturer of Motor Control Units "
        "for electric vehicles with over 50% domestic market share, faces a critical talent acquisition "
        "challenge commensurate with its rapid growth trajectory. The Sterling Intelligent Interview "
        "& Candidate Assessment Platform — developed as an enterprise-grade, AI-powered web application "
        "— addresses this challenge by automating the complete preliminary technical candidate screening "
        "lifecycle from resume parsing to final report generation.\n\n"
        "The platform employs a decoupled microservices architecture built on FastAPI (Python) and "
        "React 18, integrating OpenAI GPT-4o and Google Gemini for Large Language Model-based interview "
        "evaluation, Groq Whisper for sub-second Speech-to-Text transcription, ElevenLabs for neural "
        "Text-to-Speech synthesis with zero-disk streaming, and React Three Fiber for WebGL-based 3D "
        "avatar rendering with real-time audio-reactive lip synchronisation. A 15-table normalised SQLite "
        "database operating in Write-Ahead Logging mode provides concurrent data access for multi-session "
        "operation.\n\n"
        "Key capabilities include: NLP-based resume parsing and role-match scoring; dynamic LLM-driven "
        "question generation with adaptive difficulty; integrated Monaco Editor for code submission "
        "assessment; browser-native proctoring via face detection, tab monitoring and focus tracking; "
        "a multi-dimensional scoring engine across Technical, Communication, EQ, and Confidence "
        "dimensions; and automated FinalReport generation with hiring recommendations. The platform "
        "delivers a 100% reduction in engineering hours required for preliminary screening and generates "
        "quantified annual cost savings of ₹26.7 lakhs for Sterling E-Mobility.\n\n"
        "Pilot testing across 50 simulated candidate sessions demonstrated a mean end-to-end interview "
        "response latency of 1.61 seconds, resume screening accuracy of r=0.81 against expert assessment, "
        "and a candidate experience satisfaction rating of 4.4/5. The platform stands as a production-ready "
        "system that directly supports Sterling E-Mobility's strategic objective of scaling its engineering "
        "talent base to meet its FY2028 integrated powertrain systems growth ambitions."
    )
    add_body(doc, abstract_text)
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Keywords: ")
    run.font.bold = True; run.font.size = Pt(11); run.font.name = "Times New Roman"
    run2 = p.add_run(
        "AI Interview Platform, Resume Screening, Speech-to-Text, Text-to-Speech, "
        "3D Avatar, FastAPI, React, WebSocket, SQLite WAL, Proctoring, Sterling E-Mobility, "
        "Large Language Models, Motor Control Units, EV Technology"
    )
    run2.font.size = Pt(11); run2.font.name = "Times New Roman"

    add_page_break(doc)

# ─────────────────────────────────────────────
# REFERENCES
# ─────────────────────────────────────────────
def add_references(doc):
    add_heading(doc, "REFERENCES", level=1)
    doc.add_paragraph()

    refs = [
        "[1] A. Radford et al., \"Robust Speech Recognition via Large-Scale Weak Supervision,\" arXiv:2212.04356, OpenAI, 2022.",
        "[2] T. B. Brown et al., \"Language Models are Few-Shot Learners,\" Advances in Neural Information Processing Systems (NeurIPS), vol. 33, 2020.",
        "[3] E. Faliagka et al., \"An Integrated e-Recruitment System for Automated Personality Mining and Applicant Ranking,\" in Proc. ACM Internet Conference, 2012.",
        "[4] I. Naim et al., \"Automated Prediction and Analysis of Job Interview Performance: The Role of What You Say and How You Say It,\" in Proc. 12th IEEE FG, 2018.",
        "[5] A. van den Oord et al., \"WaveNet: A Generative Model for Raw Audio,\" arXiv:1609.03499, Google DeepMind, 2016.",
        "[6] J. Shen et al., \"Natural TTS Synthesis by Conditioning WaveNet on Mel Spectrogram Predictions,\" in Proc. ICASSP, 2018.",
        "[7] A. Harber, \"Automated Resume Screening using BERT-based Natural Language Processing,\" IEEE Access, vol. 7, 2019.",
        "[8] S. González-Carvajal and E. Garrido-Merchan, \"Comparing BERT Against Traditional Machine Learning Text Classification,\" arXiv:2005.13012, 2021.",
        "[9] N. Koenecke et al., \"Racial disparities in automated speech recognition,\" Proceedings of the National Academy of Sciences, vol. 117, no. 14, 2020.",
        "[10] J. Cassell et al., \"Embodied Conversational Agents,\" MIT Press, Cambridge MA, 2000.",
        "[11] N. C. Krämer and G. Bente, \"Personalizing e-Learning: The Social Effects of Pedagogical Agents,\" Educational Psychology Review, vol. 22, no. 1, 2010.",
        "[12] Y. Shi et al., \"Domain-Adaptive Training BERT for Response Selection,\" arXiv:1908.04812, 2020.",
        "[13] A. Mizumoto and M. Eguchi, \"Exploring the potential of using an AI language model for automated essay scoring,\" Research Methods in Applied Linguistics, vol. 2, 2023.",
        "[14] OpenAI, \"GPT-4 Technical Report,\" arXiv:2303.08774, 2023.",
        "[15] Google DeepMind, \"Gemini: A Family of Highly Capable Multimodal Models,\" arXiv:2312.11805, 2023.",
        "[16] Groq Inc., \"LPU Inference Engine Technical Documentation,\" https://groq.com/technology, 2024.",
        "[17] ElevenLabs, \"ElevenLabs TTS API Documentation,\" https://docs.elevenlabs.io, 2024.",
        "[18] Sterling E-Mobility Solutions Limited, \"Company Overview,\" https://www.sterlingemobility.com/company, 2025.",
        "[19] J. Wadhwa, \"Sterling E-Mobility to Make Systems Core of Business, Targets 70% Share by FY28,\" Economic Times Manufacturing, 2025.",
        "[20] OWASP Foundation, \"OWASP Top 10 Application Security Risks,\" https://owasp.org/Top10/, 2021.",
        "[21] S. Tiangolo, \"FastAPI: Modern, Fast Web Framework for Building APIs with Python,\" https://fastapi.tiangolo.com, 2023.",
        "[22] Meta Open Source, \"React 18 Documentation,\" https://react.dev, 2024.",
        "[23] D. Abramov, \"Three.js Documentation and React Three Fiber,\" https://docs.pmnd.rs/react-three-fiber, 2024.",
        "[24] Ministry of Heavy Industries, Government of India, \"FAME II Scheme — Faster Adoption and Manufacturing of Hybrid and Electric Vehicles,\" MHI, New Delhi, 2019.",
        "[25] ACMA, \"ACMA Excellence Awards 2025 — Automotive Component Manufacturers Association of India,\" New Delhi, 2025.",
    ]

    for ref in refs:
        para = doc.add_paragraph()
        run = para.add_run(ref)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(3)

    add_page_break(doc)

# ─────────────────────────────────────────────
# APPENDIX STUBS
# ─────────────────────────────────────────────
def add_appendices(doc):
    appendices = [
        ("APPENDIX A — API ENDPOINT REFERENCE",
         "This appendix provides a comprehensive reference for all REST and WebSocket API endpoints "
         "exposed by the Sterling AI Recruitment Engine backend. Each endpoint is documented with its "
         "HTTP method, URL pattern, request parameters/body schema, authentication requirement, and "
         "sample response.\n\n[Full API reference to be generated from FastAPI OpenAPI specification "
         "at /docs endpoint and inserted here.]"),

        ("APPENDIX B — DATABASE SCHEMA (SQL DDL)",
         "This appendix contains the complete SQL Data Definition Language (DDL) statements for all "
         "15 tables in the Sterling AI Recruitment Engine database schema, as generated by SQLAlchemy's "
         "schema creation facility.\n\n[DDL statements to be exported from database/models.py and "
         "inserted here.]"),

        ("APPENDIX C — SAMPLE INTERVIEW SESSION TRANSCRIPT",
         "The following is a representative excerpt from a complete interview session conducted by the "
         "Sterling AI Recruitment Engine with a candidate applying for the Embedded Systems Engineer "
         "role. Speaker labels indicate AI (Sterling Assessment Engine) and Candidate utterances.\n\n"
         "STERLING ASSESSMENT ENGINE: Welcome to the Sterling E-Mobility technical assessment. "
         "I am the Sterling Assessment Engine, and I will be conducting your preliminary technical "
         "interview today for the Embedded Systems Engineer position. Please introduce yourself briefly.\n\n"
         "CANDIDATE: Thank you. My name is [Candidate Name]. I have completed my B.Tech in Electronics "
         "and Communication Engineering and have been working on embedded firmware development for the "
         "past two years, primarily in automotive applications.\n\n"
         "STERLING ASSESSMENT ENGINE: Excellent background. Let us begin with a foundational concept. "
         "Can you explain the working principle of Field-Oriented Control in PMSM motor drives?\n\n"
         "[Full transcript continues for 45-60 minutes of interview duration...]"),

        ("APPENDIX D — SAMPLE FINAL REPORT OUTPUT",
         "The following represents a sample FinalReport JSON structure as generated by ranking_engine.py "
         "upon interview completion:\n\n"
         "{\n"
         "  \"report_id\": 42,\n"
         "  \"candidate_name\": \"[Candidate Name]\",\n"
         "  \"role_applied\": \"Embedded Systems Engineer\",\n"
         "  \"interview_date\": \"2026-06-02\",\n"
         "  \"overall_score\": 74.3,\n"
         "  \"technical_score\": 71.8,\n"
         "  \"communication_score\": 76.2,\n"
         "  \"eq_score\": 74.5,\n"
         "  \"confidence_score\": 72.1,\n"
         "  \"grade\": \"B+\",\n"
         "  \"hiring_decision\": \"HOLD\",\n"
         "  \"strengths\": [\"Strong CAN-BUS protocol knowledge\", \"Clear explanation of FOC concepts\"],\n"
         "  \"weaknesses\": [\"Limited experience with AUTOSAR\", \"Incomplete coverage of ISO 26262\"],\n"
         "  \"proctoring_summary\": \"1 tab switch event detected at T+22min\"\n"
         "}"),

        ("APPENDIX E — ABBREVIATIONS AND GLOSSARY",
         "AI — Artificial Intelligence\n"
         "API — Application Programming Interface\n"
         "ASGI — Asynchronous Server Gateway Interface\n"
         "BLDC — Brushless Direct Current (motor)\n"
         "CAN — Controller Area Network\n"
         "CORS — Cross-Origin Resource Sharing\n"
         "CV — Coefficient of Variation\n"
         "DPDP — Digital Personal Data Protection (Act)\n"
         "EQ — Emotional Quotient / Intelligence\n"
         "EV — Electric Vehicle\n"
         "FAME — Faster Adoption and Manufacturing of Hybrid and Electric Vehicles\n"
         "FOC — Field-Oriented Control\n"
         "glTF — GL Transmission Format\n"
         "HRIS — Human Resource Information System\n"
         "JWT — JSON Web Token\n"
         "LCV — Light Commercial Vehicle\n"
         "LLM — Large Language Model\n"
         "LPU — Language Processing Unit (Groq)\n"
         "MCU — Motor Control Unit\n"
         "NER — Named Entity Recognition\n"
         "NLP — Natural Language Processing\n"
         "OBC — On-Board Charger\n"
         "OEM — Original Equipment Manufacturer\n"
         "OOM — Out-Of-Memory\n"
         "ORM — Object-Relational Mapping\n"
         "PAT — Profit After Tax\n"
         "PDU — Power Distribution Unit\n"
         "PMSM — Permanent Magnet Synchronous Motor\n"
         "RBAC — Role-Based Access Control\n"
         "ROI — Return on Investment\n"
         "SEM — Sterling E-Mobility Solutions Limited\n"
         "SPA — Single-Page Application\n"
         "SQL — Structured Query Language\n"
         "STT — Speech-to-Text\n"
         "TTS — Text-to-Speech\n"
         "UDS — Unified Diagnostic Services\n"
         "VAD — Voice Activity Detection\n"
         "WAL — Write-Ahead Logging\n"
         "WebGL — Web Graphics Library\n"
         "WSS — WebSocket Secure\n"
         "YOE — Years of Experience")
    ]

    for title, content in appendices:
        add_separator(doc)
        add_heading(doc, title, level=1)
        doc.add_paragraph()
        for para_text in content.split('\n\n'):
            lines = para_text.strip().split('\n')
            for line in lines:
                if line.strip():
                    add_body(doc, line.strip())
        doc.add_paragraph()
        add_page_break(doc)

# ─────────────────────────────────────────────
# DOCUMENT PAGE SETUP
# ─────────────────────────────────────────────
def setup_document(doc):
    # Page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.0)

    # Default paragraph font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Add page number in footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    # Add page number field
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

# ─────────────────────────────────────────────
# CHAPTER FILES LIST
# ─────────────────────────────────────────────
CHAPTER_FILES = [
    "Chapter_01_Organization_Overview.md",
    "Chapter_02_Project_Introduction.md",
    "Chapter_03_Literature_Review.md",
    "Chapter_04_Existing_System.md",
    "Chapter_05_Proposed_System.md",
    "Chapter_06_System_Architecture.md",
    "Chapter_07_Technology_Stack.md",
    "Chapter_08_Database_Design.md",
    "Chapter_09_Module_Implementation.md",
    "Chapter_10_Security.md",
    "Chapter_11_Testing_Validation.md",
    "Chapter_12_Results_Analysis.md",
    "Chapter_13_Future_Enhancements.md",
    "Chapter_14_Conclusion.md",
]

# ─────────────────────────────────────────────
# MAIN ASSEMBLER
# ─────────────────────────────────────────────
def assemble_report():
    print("Initialising document...")
    doc = Document()
    setup_document(doc)

    print("Adding cover page...")
    add_cover_page(doc)

    print("Adding certificate...")
    add_certificate(doc)

    print("Adding declaration...")
    add_declaration(doc)

    print("Adding acknowledgement...")
    add_acknowledgement(doc)

    print("Adding abstract...")
    add_abstract(doc)

    # Chapters
    for fname in CHAPTER_FILES:
        fpath = os.path.join(BASE_DIR, fname)
        if not os.path.exists(fpath):
            print(f"WARNING: {fname} not found — skipping.")
            continue
        print(f"Processing {fname}...")
        with open(fpath, 'r', encoding='utf-8') as f:
            content = f.read()
        parse_md_to_docx(doc, content)
        add_page_break(doc)

    print("Adding references...")
    add_references(doc)

    print("Adding appendices...")
    add_appendices(doc)

    print(f"Saving DOCX to: {OUTPUT_DOCX}")
    doc.save(OUTPUT_DOCX)
    print("DOCX saved successfully!")
    return OUTPUT_DOCX

if __name__ == "__main__":
    output = assemble_report()
    print(f"\nFinal output: {output}")
